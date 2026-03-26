import streamlit as st
import google.generativeai as genai
import pdfplumber
import re
from docx import Document
from io import BytesIO

# --- 1. SETUP & BRANDING ---
st.set_page_config(page_title="Resume Signal Auditor", layout="wide")

st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    .stTextArea textarea { font-family: 'Courier New', Courier, monospace; font-size: 14px; line-height: 1.6; border: 1px solid #d1d1d1; }
    .stMarkdown { font-size: 15px; }
    .rewrite-highlight { color: #007bff; font-weight: bold; }
    .action-highlight { color: #dc3545; font-weight: bold; }
    </style>
    """, unsafe_allow_html=True)

if "GOOGLE_API_KEY" in st.secrets:
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
else:
    st.error("Missing API Key. Please add 'GOOGLE_API_KEY' to your Streamlit Secrets.")

# --- 2. THE AI BRAIN (Advanced Partitioning) ---
def run_ai_audit(resume_text, jd_text):
    models_to_try = ['gemini-2.0-flash', 'gemini-1.5-flash', 'gemini-flash-latest']
    
    prompt = f"""
    You are a Staff-level Hiring Manager. Audit this resume against the Job Description.
    
    YOUR TASK:
    1. CHANGELOG: List 3-4 specific edits you made to the tone or vocabulary (e.g., "Changed 'Manager' to 'Orchestrator' to match JD's focus on leadership").
    2. THE DRAFT: Rewrite the resume to mirror the JD's language.
       - Use {{REWRITE: phrase}} whenever you significantly change a phrase for tone or keywords.
       - Use [ACTION: instruction] whenever you identify a missing metric or fact.
    
    RESUME: {resume_text}
    JD: {jd_text}

    OUTPUT FORMAT:
    CHANGELOG: [List specific pivots here]
    KEYWORDS: [8 keywords]
    DRAFT: [Full resume with {{REWRITE: ...}} and [ACTION: ...]]
    """
    
    for model_name in models_to_try:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            return response.text
        except Exception:
            continue
    return None

# --- 3. HELPER TOOLS ---
def create_docx(text):
    doc = Document()
    # Remove all AI helper tags for the final document
    clean_text = re.sub(r'\{(?:REWRITE): (.*?)\}', r'\1', text) 
    clean_text = re.sub(r'\[(?:ACTION|METRIC|QUERY):?.*?\]', '', clean_text, flags=re.IGNORECASE) 
    
    for line in clean_text.split('\n'):
        if line.strip():
            p = doc.add_paragraph()
            if "|" in line or line.isupper():
                p.add_run(line).bold = True
            else:
                p.add_run(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 4. THE UI INTERFACE ---
st.title("🎯 Resume Signal Auditor")
st.caption("Beta v1.1.4 | Differentiating AI Edits vs. User Actions")

with st.sidebar:
    st.header("1. Input Data")
    uploaded_file = st.file_uploader("Upload Resume (PDF)", type="pdf")
    jd_input = st.text_area("Target Job Description", height=250)
    
    if st.button("Run Strategic Audit", use_container_width=True):
        if uploaded_file and jd_input:
            with st.spinner("Analyzing signals..."):
                with pdfplumber.open(uploaded_file) as pdf:
                    resume_text = "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])
                
                output = run_ai_audit(resume_text, jd_input)
                if output and "DRAFT:" in output:
                    st.session_state['changelog'] = output.split("CHANGELOG:")[1].split("KEYWORDS:")[0].strip()
                    st.session_state['keywords'] = output.split("KEYWORDS:")[1].split("DRAFT:")[0].strip()
                    st.session_state['draft'] = output.split("DRAFT:")[1].strip()
                    # Capture Action items only
                    st.session_state['actions'] = re.findall(r'\[ACTION:? (.*?)\]', st.session_state['draft'])
                else:
                    st.error("AI Error. Try again.")

# --- 5. THE AUDIT DASHBOARD ---
if 'draft' in st.session_state:
    
    # NEW: Header section showing the Rationale and Changelog
    with st.expander("🛠️ AI Changelog: See what I changed for you", expanded=True):
        st.info(st.session_state['changelog'])
        st.write(f"**Target Keywords Integrated:** {st.session_state['keywords']}")

    col_ed, col_check = st.columns([1.8, 1], gap="large")

    with col_ed:
        st.subheader("✍️ Editor")
        st.caption("Instructions: 1) Review {REWRITE} items (Blue). 2) Resolve [ACTION] items (Red).")
        
        # We display a preview of the REWRITES above the editor for readability
        preview_text = st.session_state['draft']
        preview_text = preview_text.replace("{REWRITE:", "📘 **").replace("}", "**")
        preview_text = preview_text.replace("[ACTION:", "🔴 **[ACTION:**").replace("]", "**]")
        
        with st.container(border=True):
            st.markdown(preview_text)
        
        st.divider()
        # The user edits the plain text here
        st.session_state['draft'] = st.text_area("Live Editor (Make changes here)", value=st.session_state['draft'], height=500)

    with col_check:
        st.subheader("📋 Action Items")
        st.write("Address these in the editor to unlock download.")
        
        current_tags = re.findall(r'\[ACTION:? (.*?)\]', st.session_state['draft'])
        
        resolved_count = 0
        total_actions = len(st.session_state['actions'])
        
        if total_actions > 0:
            for item in st.session_state['actions']:
                if item in current_tags:
                    st.error(f"👉 {item}")
                else:
                    st.success(f"✅ Resolved")
                    resolved_count += 1
            
            st.progress(resolved_count / total_actions)
        else:
            st.info("No specific actions found. Review the rewrites for accuracy.")

        st.divider()
        if resolved_count == total_actions:
            st.success("Ready for Download!")
            st.download_button("📥 Download Final Word Doc", data=create_docx(st.session_state['draft']), file_name="Optimized_Resume.docx", use_container_width=True)
        else:
            st.button("Download Locked", disabled=True, use_container_width=True)

else:
    st.info("👋 Upload your resume and the Job Description in the sidebar to begin.")
